from fastapi import APIRouter, Depends, HTTPException, BackgroundTasks
from sqlalchemy.orm import Session
from typing import List, Optional, Any, Dict, Literal
from datetime import datetime
from decimal import Decimal
from pydantic import BaseModel, ConfigDict, Field
from enum import Enum
from app.models import Project, ResearchStep
from app.services import ResearchService, MLXService, PerplexityService
from app.utils.database import get_db
from app.utils.security import get_current_user
import logging
from dotenv import load_dotenv
import json
import os
from datetime import timezone


# Status Enums for type safety
class ProjectStatus(str, Enum):
    NOT_STARTED = "not_started"
    IN_PROGRESS = "in_progress"
    COMPLETED = "completed"

class EvaluationStatus(str, Enum):
    PENDING = "pending"
    IN_PROGRESS = "in_progress"
    COMPLETED = "completed"
    FAILED = "failed"

class StepStatus(str, Enum):
    PENDING = "pending"
    IN_PROGRESS = "in_progress"
    COMPLETED = "completed"
    FAILED = "failed"

class StepResponse(BaseModel):
    model_config = ConfigDict(from_attributes=True)
    id: int
    project_id: int
    step_number: int
    description: str = ""
    keywords: List[str] = Field(default_factory=list)
    methodology: str = ""
    output_format: str = ""
    status: StepStatus = StepStatus.PENDING
    result: Dict[str, Any] = Field(default_factory=dict)
    executed_at: Optional[datetime] = None
    started_at: Optional[datetime] = None
    completed_at: Optional[datetime] = None
    error_message: Optional[str] = None
    progress_percentage: int = 0


class StepUpdateRequest(BaseModel):
    description: str
    keywords: List[str] = Field(default_factory=list)
    methodology: str = ""
    output_format: str = ""

# Base Models with strict typing
class ProjectCreate(BaseModel):
    title: str
    description: str
    evaluation_plan: str
    submission_format: str
    metadata1: Dict[str, Any] = Field(default_factory=dict)

class ProjectResponse(BaseModel):
    model_config = ConfigDict(from_attributes=True)
    id: int
    title: str
    description: Optional[str] = None
    evaluation_plan: Optional[str] = None
    submission_format: Optional[str] = None
    metadata1: Dict[str, Any] = Field(default_factory=dict)
    created_at: datetime
    updated_at: datetime
    user_id: int
    research_field: Optional[str] = None
    evaluation_status: EvaluationStatus = EvaluationStatus.PENDING
    final_score: Optional[Decimal] = None
    submission_date: Optional[datetime] = None
    total_steps: int = 0
    completed_steps: int = 0
    status: ProjectStatus = ProjectStatus.NOT_STARTED

class StepStatusResponse(BaseModel):
    step_number: int
    status: StepStatus
    progress: float = 0.0
    started_at: Optional[datetime] = None
    completed_at: Optional[datetime] = None
    error_message: Optional[str] = None

class ResearchStatusResponse(BaseModel):
    project_id: int
    status: EvaluationStatus
    total_steps: int
    completed_steps: int
    steps: List[StepStatusResponse]

class ResearchProgressResponse(BaseModel):
    total_steps: int
    completed_steps: int
    progress_percentage: float
    status: EvaluationStatus
    steps: List[StepStatusResponse]

class StepExecutionResponse(BaseModel):
    status: Literal["success", "failure"]
    step_number: int
    result: Dict[str, Any] = Field(default_factory=dict)

class ExportResultItem(BaseModel):
    step: str
    title: str
    status: StepStatus
    started_at: Optional[datetime]
    completed_at: Optional[datetime]
    result: Optional[str]
    final_report_en: Optional[str] = None  # 영문 최종 보고서
    final_report_kr: Optional[str] = None  # 한글 최종 보고서

# 최종 리포트를 포함한 전체 결과 응답 모델
class ResearchExportResponse(BaseModel):
    project_info: ProjectResponse
    steps: List[ExportResultItem]
    final_report_en: Optional[str] = None
    final_report_kr: Optional[str] = None

# Router setup
load_dotenv()
pplx_api_key = os.getenv("PERPLEXITY_API_KEY")
mlx_service = MLXService()
perplexityservice = PerplexityService(api_key=pplx_api_key)
logger = logging.getLogger(__name__)
router = APIRouter()

def get_research_service(db: Session = Depends(get_db)) -> ResearchService:
    try:
        mlx_service = MLXService()
        perplexity_api_key = os.getenv("PERPLEXITY_API_KEY")
        if not perplexity_api_key:
            raise HTTPException(
                status_code=500,
                detail="PERPLEXITY_API_KEY environment variable is not set"
            )
        perplexity_service = PerplexityService(api_key=perplexity_api_key)
        return ResearchService(
            db=db,
            mlx_service=mlx_service,
            perplexity_service=perplexity_service
        )
    except Exception as e:
        logger.error(f"Error creating ResearchService: {str(e)}")
        raise HTTPException(
            status_code=500,
            detail=f"Service initialization failed: {str(e)}"
        )
    
def get_project_status(project: Project) -> ProjectStatus:
    """Helper function to determine project status"""
    if project.completed_steps == project.total_steps and project.total_steps > 0:
        return ProjectStatus.COMPLETED
    elif project.completed_steps > 0:
        return ProjectStatus.IN_PROGRESS
    return ProjectStatus.NOT_STARTED

@router.put("/{project_id}/steps/{step_number}", response_model=StepResponse)
def update_step(
    project_id: int,
    step_number: int,
    step_update: StepUpdateRequest,
    db: Session = Depends(get_db),
    current_user: int = Depends(get_current_user)
) -> StepResponse:
    project = db.query(Project).filter(
        Project.id == project_id,
        Project.user_id == current_user
    ).first()
    
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    
    step = db.query(ResearchStep).filter(
        ResearchStep.project_id == project_id,
        ResearchStep.step_number == step_number
    ).first()
    
    if not step:
        raise HTTPException(status_code=404, detail="Step not found")
    
    step.description = step_update.description
    step.keywords_list = step_update.keywords  # property setter 사용
    step.methodology = step_update.methodology
    step.output_format = step_update.output_format
    
    db.commit()
    db.refresh(step)
    
    return StepResponse(
        id=step.id,
        project_id=step.project_id,
        step_number=step.step_number,
        status=StepStatus(step.status),
        description=step.description,
        keywords=step.keywords_list,  # property getter 사용
        methodology=step.methodology,
        output_format=step.output_format,
        result=step.result or {},
        executed_at=step.executed_at,
        started_at=step.started_at,
        completed_at=step.completed_at,
        error_message=step.error_message,
        progress_percentage=step.progress_percentage or 0
    )

@router.get("/{project_id}/steps", response_model=List[StepResponse])
def list_project_steps(
    project_id: int,
    db: Session = Depends(get_db),
    current_user: int = Depends(get_current_user)
) -> List[StepResponse]:
    project = db.query(Project).filter(
        Project.id == project_id,
        Project.user_id == current_user
    ).first()
    
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")

    steps = db.query(ResearchStep).filter(ResearchStep.project_id == project_id).all()
    
    return [
        StepResponse(
            id=step.id,
            project_id=step.project_id,
            step_number=step.step_number,
            status=StepStatus(step.status),
            description=step.description or "",
            keywords=step.keywords_list,  # property getter 사용
            methodology=step.methodology or "",
            output_format=step.output_format or "",
            result=step.result or {},
            executed_at=step.executed_at,
            started_at=step.started_at,
            completed_at=step.completed_at,
            error_message=step.error_message,
            progress_percentage=step.progress_percentage or 0
        )
        for step in steps
    ]

@router.get("/{project_id}/steps/{step_number}", response_model=StepResponse)
def get_project_step(
    project_id: int,
    step_number: int,
    db: Session = Depends(get_db),
    current_user: int = Depends(get_current_user)
) -> StepResponse:
    project = db.query(Project).filter(
        Project.id == project_id,
        Project.user_id == current_user
    ).first()
    
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    
    step = db.query(ResearchStep).filter(
        ResearchStep.project_id == project_id,
        ResearchStep.step_number == step_number
    ).first()
    
    if not step:
        raise HTTPException(status_code=404, detail="Step not found")
    
    return StepResponse(
        id=step.id,
        project_id=step.project_id,
        step_number=step.step_number,
        description=step.description or "",
        keywords=step.keywords_list,  # property getter 사용
        methodology=step.methodology or "",
        output_format=step.output_format or "",
        status=StepStatus(step.status),
        result=step.result or {},
        executed_at=step.executed_at,
        started_at=step.started_at,
        completed_at=step.completed_at,
        error_message=step.error_message,
        progress_percentage=step.progress_percentage or 0
    )

@router.post("/", response_model=ProjectResponse)
def create_project(
    project: ProjectCreate,
    db: Session = Depends(get_db),
    current_user: int = Depends(get_current_user)
) -> ProjectResponse:
    try:
        # MLXService와 PerplexityService 초기화
        mlx_service = MLXService()
        
        # Perplexity API 키 확인
        perplexity_api_key = os.getenv("PERPLEXITY_API_KEY")
        if not perplexity_api_key:
            raise HTTPException(
                status_code=500,
                detail="PERPLEXITY_API_KEY environment variable is not set"
            )
        
        # PerplexityService 인스턴스 생성
        perplexity_service = PerplexityService(api_key=perplexity_api_key)
        
        # ResearchService 인스턴스 생성 - 필요한 의존성 모두 전달
        research_service = ResearchService(
            db=db,
            mlx_service=mlx_service,
            perplexity_service=perplexity_service
        )
        
        # 프로젝트 생성
        new_project = research_service.create_project(
            user_id=current_user,
            project_data=project.model_dump()
        )
        
        return ProjectResponse(
            id=new_project.id,
            title=new_project.title,
            description=new_project.description,
            evaluation_plan=new_project.evaluation_plan,
            submission_format=new_project.submission_format,
            metadata1=new_project.metadata1 or {},
            created_at=new_project.created_at,
            updated_at=new_project.updated_at,
            user_id=current_user,
            status=ProjectStatus.NOT_STARTED,
            evaluation_status=EvaluationStatus.PENDING
        )
    
    except Exception as e:
        logger.error(f"Error creating project: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

@router.get("/", response_model=List[ProjectResponse])
def list_projects(
    db: Session = Depends(get_db),
    current_user: int = Depends(get_current_user)
) -> List[ProjectResponse]:
    projects = db.query(Project).filter(Project.user_id == current_user).all()
    return [
        ProjectResponse(
            id=project.id,
            title=project.title,
            description=project.description,
            evaluation_plan=project.evaluation_plan,
            submission_format=project.submission_format,
            metadata1=project.metadata1 or {},
            created_at=project.created_at,
            updated_at=project.updated_at,
            user_id=project.user_id,
            research_field=project.research_field,
            evaluation_status=EvaluationStatus(project.evaluation_status),
            final_score=project.final_score,
            submission_date=project.submission_date,
            total_steps=project.total_steps,
            completed_steps=project.completed_steps,
            status=get_project_status(project)
        )
        for project in projects
    ]

@router.get("/{project_id}", response_model=ProjectResponse)
def get_project(
    project_id: int,
    db: Session = Depends(get_db),
    current_user: int = Depends(get_current_user)
) -> ProjectResponse:
    project = db.query(Project).filter(
        Project.id == project_id,
        Project.user_id == current_user
    ).first()
    
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    
    return ProjectResponse(
        id=project.id,
        title=project.title,
        description=project.description,
        evaluation_plan=project.evaluation_plan,
        submission_format=project.submission_format,
        metadata1=project.metadata1 or {},
        created_at=project.created_at,
        updated_at=project.updated_at,
        user_id=project.user_id,
        research_field=project.research_field,
        evaluation_status=EvaluationStatus(project.evaluation_status),
        final_score=project.final_score,
        submission_date=project.submission_date,
        total_steps=project.total_steps,
        completed_steps=project.completed_steps,
        status=get_project_status(project)
    )

@router.put("/{project_id}", response_model=ProjectResponse)
def update_project(
    project_id: int,
    project_update: ProjectCreate,
    db: Session = Depends(get_db),
    current_user: int = Depends(get_current_user)
) -> ProjectResponse:
    project = db.query(Project).filter(
        Project.id == project_id,
        Project.user_id == current_user
    ).first()
    
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    
    project.title = project_update.title
    project.description = project_update.description
    project.evaluation_plan = project_update.evaluation_plan
    project.submission_format = project_update.submission_format
    project.metadata1 = project_update.metadata1
    project.updated_at = datetime.now(timezone.utc)
    
    db.commit()
    db.refresh(project)
    
    return ProjectResponse(
        id=project.id,
        title=project.title,
        description=project.description,
        evaluation_plan=project.evaluation_plan,
        submission_format=project.submission_format,
        metadata1=project.metadata1,
        created_at=project.created_at,
        updated_at=project.updated_at,
        user_id=project.user_id,
        status=get_project_status(project)
    )

@router.delete("/{project_id}")
def delete_project(
    project_id: int,
    db: Session = Depends(get_db),
    current_user: int = Depends(get_current_user)
) -> Dict[str, str]:
    project = db.query(Project).filter(
        Project.id == project_id,
        Project.user_id == current_user
    ).first()
    
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    
    db.query(ResearchStep).filter(ResearchStep.project_id == project_id).delete()
    db.delete(project)
    db.commit()
    
    return {"message": "Project deleted successfully"}


@router.get("/{project_id}/steps/{step_number}/result")
def get_step_result(
    project_id: int,
    step_number: int,
    db: Session = Depends(get_db),
    current_user: int = Depends(get_current_user)
) -> Dict[str, Any]:
    project = db.query(Project).filter(
        Project.id == project_id,
        Project.user_id == current_user
    ).first()
    
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    
    step = db.query(ResearchStep).filter(
        ResearchStep.project_id == project_id,
        ResearchStep.step_number == step_number
    ).first()
    
    if not step:
        raise HTTPException(status_code=404, detail="Step not found")
    
    if not step.result:
        raise HTTPException(status_code=404, detail="No result found for this step")
    
    return {
        "step_id": step.id,
        "step_number": step.step_number,
        "status": step.status,
        "result": step.result,
        "executed_at": step.executed_at.isoformat() if step.executed_at else None
    }

@router.post("/{project_id}/steps/reset")
def reset_project_steps(
    project_id: int,
    db: Session = Depends(get_db),
    current_user: int = Depends(get_current_user)
) -> Dict[str, str]:
    project = db.query(Project).filter(
        Project.id == project_id,
        Project.user_id == current_user
    ).first()
    
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    
    steps = db.query(ResearchStep).filter(ResearchStep.project_id == project_id).all()
    for step in steps:
        step.status = StepStatus.PENDING.value
        step.result = None
        step.executed_at = None
    
    db.commit()
    
    return {"message": "All steps have been reset successfully"}



@router.post("/{project_id}/steps/{step_number}/execute")
async def execute_step(
    project_id: int,
    step_number: int,
    db: Session = Depends(get_db),
    current_user: int = Depends(get_current_user)
) -> StepExecutionResponse:
    project = db.query(Project).filter(
        Project.id == project_id,
        Project.user_id == current_user
    ).first()
    
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
        
    step = db.query(ResearchStep).filter(
        ResearchStep.project_id == project_id,
        ResearchStep.step_number == step_number
    ).first()
    
    if not step:
        raise HTTPException(status_code=404, detail="Research step not found")

    try:
        research_service = ResearchService(
            db=db,
            mlx_service=MLXService(),
            perplexity_service=PerplexityService(os.getenv("PERPLEXITY_API_KEY"))
        )
        
        result = research_service.execute_step(project_id, step_number)
        
        step.status = StepStatus.COMPLETED.value
        step.result = result
        step.executed_at = datetime.now(timezone.utc)
        db.commit()
        
        return StepExecutionResponse(
            status="success",
            step_number=step_number,
            result=result
        )
        
    except Exception as e:
        logger.error(f"Step execution failed: {str(e)}")
        step.status = StepStatus.FAILED.value
        db.commit()
        raise HTTPException(status_code=500, detail=f"Step execution failed: {str(e)}")

@router.post("/{project_id}/execute")
async def execute_research(
    project_id: int,
    background_tasks: BackgroundTasks,
    db: Session = Depends(get_db),
    current_user: int = Depends(get_current_user)
) -> Dict[str, Any]:
    try:
        project = db.query(Project).filter(
            Project.id == project_id,
            Project.user_id == current_user
        ).first()
        
        if not project:
            raise HTTPException(status_code=404, detail="Project not found")
            
        if project.evaluation_status == EvaluationStatus.IN_PROGRESS.value:
            raise HTTPException(status_code=400, detail="Research is already in progress")
            
        steps = db.query(ResearchStep).filter(
            ResearchStep.project_id == project_id
        ).order_by(ResearchStep.step_number).all()
        
        if not steps:
            raise HTTPException(status_code=400, detail="No research steps defined for this project")

        perplexity_api_key = os.getenv("PERPLEXITY_API_KEY")
        
        if not perplexity_api_key:
            raise HTTPException(status_code=500, detail="API keys not properly configured")
            
        research_service = ResearchService(
            db=db,
            mlx_service=MLXService(),
            perplexity_service=PerplexityService(perplexity_api_key)
        )
        
        # Update project status
        project.evaluation_status = EvaluationStatus.IN_PROGRESS.value
        project.total_steps = len(steps)
        project.completed_steps = 0
        db.commit()
        
        # Reset all steps
        for step in steps:
            step.status = StepStatus.PENDING.value
            step.result = None
            step.started_at = None
            step.completed_at = None
            step.error_message = None
            step.progress_percentage = 0
        db.commit()

        async def execute_research_steps():
            try:
                await research_service.execute_all_steps(project_id)
            except Exception as e:
                logger.error(f"Error executing research steps for project {project_id}: {str(e)}")
                with db.begin():
                    project = db.query(Project).filter(Project.id == project_id).first()
                    if project:
                        project.evaluation_status = EvaluationStatus.FAILED.value
                        db.commit()

        background_tasks.add_task(execute_research_steps)
        
        return {
            "status": "started",
            "message": "Research has been started",
            "project_id": project_id,
            "total_steps": len(steps)
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error starting research for project {project_id}: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

@router.get("/{project_id}/status", response_model=ResearchStatusResponse)
def get_research_status(
    project_id: int,
    db: Session = Depends(get_db),
    current_user: int = Depends(get_current_user)
) -> ResearchStatusResponse:
    try:
        project = db.query(Project).filter(
            Project.id == project_id,
            Project.user_id == current_user
        ).first()
        
        if not project:
            raise HTTPException(status_code=404, detail="Project not found")
            
        steps = db.query(ResearchStep).filter(
            ResearchStep.project_id == project_id
        ).order_by(ResearchStep.step_number).all()
        
        step_statuses = [
            StepStatusResponse(
                step_number=step.step_number,
                status=StepStatus(step.status),
                progress=step.progress_percentage or 0.0,
                started_at=step.started_at,
                completed_at=step.completed_at,
                error_message=step.error_message
            )
            for step in steps
        ]
        
        return ResearchStatusResponse(
            project_id=project_id,
            status=EvaluationStatus(project.evaluation_status),
            total_steps=project.total_steps,
            completed_steps=project.completed_steps,
            steps=step_statuses
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error getting research status for project {project_id}: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))
@router.get("/{project_id}/export", response_model=List[ExportResultItem])
def export_results(
    project_id: int,
    db: Session = Depends(get_db),
    current_user: int = Depends(get_current_user)
) -> List[ExportResultItem]:
    try:
        project = db.query(Project).filter(
            Project.id == project_id,
            Project.user_id == current_user
        ).first()
        
        if not project:
            raise HTTPException(status_code=404, detail="Project not found")
        
        steps = db.query(ResearchStep).filter(
            ResearchStep.project_id == project_id
        ).order_by(ResearchStep.step_number).all()
        
        results = []
        for step in steps:
            try:
                # result가 문자열인 경우 JSON으로 파싱 시도
                result_content = step.result
                if isinstance(step.result, str):
                    try:
                        result_content = json.loads(step.result)
                    except json.JSONDecodeError:
                        result_content = {"content": step.result}
                elif not isinstance(step.result, dict):
                    result_content = {"content": str(step.result)}
                
                results.append(
                    ExportResultItem(
                        step=f"Step {step.step_number}",
                        title=step.description or "",
                        status=StepStatus(step.status),
                        started_at=step.started_at,
                        completed_at=step.completed_at,
                        result=json.dumps(result_content, ensure_ascii=False)
                    )
                )
            except Exception as e:
                logger.error(f"Error processing step {step.step_number}: {str(e)}")
                results.append(
                    ExportResultItem(
                        step=f"Step {step.step_number}",
                        title=step.description or "",
                        status=StepStatus(step.status),
                        started_at=step.started_at,
                        completed_at=step.completed_at,
                        result=json.dumps({"error": str(e)}, ensure_ascii=False)
                    )
                )
        
        # 최종 보고서 정보 추가
        if project.metadata1 and "final_reports" in project.metadata1:
            final_reports = project.metadata1["final_reports"]
            results.append(
                ExportResultItem(
                    step="Final Report",
                    title="최종 연구 보고서 (영문)",
                    status=StepStatus.COMPLETED,
                    started_at=None,
                    completed_at=datetime.fromisoformat(final_reports.get("generated_at", "")) if final_reports.get("generated_at") else None,
                    result=json.dumps({"content": final_reports.get("final_report_en", "")}, ensure_ascii=False)
                )
            )
            results.append(
                ExportResultItem(
                    step="Final Report KR",
                    title="최종 연구 보고서 (한글)",
                    status=StepStatus.COMPLETED,
                    started_at=None,
                    completed_at=datetime.fromisoformat(final_reports.get("generated_at", "")) if final_reports.get("generated_at") else None,
                    result=json.dumps({"content": final_reports.get("final_report_kr", "")}, ensure_ascii=False)
                )
            )
        
        return results
    
    except Exception as e:
        logger.error(f"Export failed: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Export failed: {str(e)}")

@router.get("/{project_id}/export/docx")
async def export_results_docx(
    project_id: int,
    db: Session = Depends(get_db),
    current_user: int = Depends(get_current_user)
):
    try:
        from docx import Document
        from docx.shared import Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from io import BytesIO
        
        # 결과 데이터 가져오기
        export_items = export_results(project_id, db, current_user)
        project = db.query(Project).get(project_id)
        
        # DOCX 문서 생성
        doc = Document()
        
        # 제목
        title = doc.add_heading(project.title, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 프로젝트 정보
        doc.add_heading('프로젝트 개요', level=1)
        doc.add_paragraph(f"설명: {project.description}")
        doc.add_paragraph(f"평가 계획: {project.evaluation_plan}")
        doc.add_paragraph(f"제출 형식: {project.submission_format}")
        
        # 연구 단계별 결과
        doc.add_heading('연구 단계별 결과', level=1)
        for item in export_items:
            if item.step.startswith("Final Report"):
                continue  # 최종 보고서는 나중에 별도로 처리
                
            doc.add_heading(f"{item.step}: {item.title}", level=2)
            
            # 상태 정보
            p = doc.add_paragraph()
            p.add_run('상태: ').bold = True
            p.add_run(str(item.status))
            
            if item.started_at:
                p = doc.add_paragraph()
                p.add_run('시작 시간: ').bold = True
                p.add_run(item.started_at.strftime("%Y-%m-%d %H:%M:%S"))
            
            if item.completed_at:
                p = doc.add_paragraph()
                p.add_run('완료 시간: ').bold = True
                p.add_run(item.completed_at.strftime("%Y-%m-%d %H:%M:%S"))
            
            # 결과 내용
            if item.result:
                doc.add_heading('결과', level=3)
                try:
                    result_data = json.loads(item.result)
                    if isinstance(result_data, dict):
                        for key, value in result_data.items():
                            p = doc.add_paragraph()
                            if key != "content":
                                p.add_run(f"{key}: ").bold = True
                            p.add_run(str(value))
                    else:
                        doc.add_paragraph(str(result_data))
                except json.JSONDecodeError:
                    doc.add_paragraph(item.result)
            
            doc.add_paragraph('_' * 50)
        
        # 최종 보고서
        final_reports = [item for item in export_items if item.step.startswith("Final Report")]
        if final_reports:
            for report in final_reports:
                doc.add_heading(report.title, level=1)
                if report.result:
                    try:
                        result_data = json.loads(report.result)
                        doc.add_paragraph(result_data.get("content", ""))
                    except json.JSONDecodeError:
                        doc.add_paragraph(report.result)
        
        # 메모리에 문서 저장
        file_stream = BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        
        # 파일 이름 설정
        filename = f"연구결과_{project_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        
        headers = {
            'Content-Disposition': f'attachment; filename="{filename}"',
            'Content-Type': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        }
        
        return Response(
            content=file_stream.getvalue(),
            headers=headers,
            media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
    except Exception as e:
        logger.error(f"DOCX export failed: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

@router.get("/{project_id}/progress", response_model=ResearchProgressResponse)
def get_progress(
    project_id: int,
    db: Session = Depends(get_db),
    current_user: int = Depends(get_current_user)
) -> ResearchProgressResponse:
    project = db.query(Project).filter(
        Project.id == project_id,
        Project.user_id == current_user
    ).first()
    
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    
    steps = db.query(ResearchStep).filter(
        ResearchStep.project_id == project_id
    ).order_by(ResearchStep.step_number).all()
    
    total_steps = len(steps)
    completed_steps = sum(1 for step in steps if step.status == StepStatus.COMPLETED.value)
    
    step_statuses = [
        StepStatusResponse(
            step_number=step.step_number,
            status=StepStatus(step.status),
            progress=step.progress_percentage or 0.0,
            started_at=step.started_at,
            completed_at=step.completed_at,
            error_message=step.error_message
        )
        for step in steps
    ]
    
    return ResearchProgressResponse(
        total_steps=total_steps,
        completed_steps=completed_steps,
        progress_percentage=(completed_steps / total_steps * 100) if total_steps > 0 else 0,
        status=EvaluationStatus(project.evaluation_status),
     
        steps=step_statuses
    )