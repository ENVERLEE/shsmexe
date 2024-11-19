from typing import List, Dict, Optional
from datetime import datetime
import pandas as pd
from docx import Document
from docx.shared import Pt
import numpy as np  # 누락된 import 추가
import json
from sqlalchemy.orm import Session
from app.models import Project, ResearchStep
import logging
from fastapi import HTTPException
import anthropic 
import voyageai
from dotenv import load_dotenv
from .mlx_service import MLXService
from .perplexity_service import PerplexityService, SearchConfig
import os
import re
from concurrent.futures import ThreadPoolExecutor
import requests
import ssl
import aiohttp
import asyncio

load_dotenv()
XAI_API_KEY = os.getenv("XAI_API_KEY")

API_URL = "https://api-inference.huggingface.co/models/sentence-transformers/paraphrase-multilingual-MiniLM-L12-v2"
headers = {"Authorization": "Bearer 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"}  # Replace with your actual token
class ResearchService:
    def __init__(self, db: Session, mlx_service: MLXService, perplexity_service: PerplexityService):
        """Initialize ResearchService with improved SSL handling"""
        self.db = db
        self.mlx_service = mlx_service
        
        # PerplexityService 초기화 개선
        try:
            perplexity_api_key = os.getenv("PERPLEXITY_API_KEY")
            if not perplexity_api_key:
                raise ValueError("PERPLEXITY_API_KEY environment variable not found")
            
            self.perplexity_service = PerplexityService(api_key=perplexity_api_key)
            self.logger = logging.getLogger(__name__)
            
        except Exception as e:
            self.logger.error(f"Failed to initialize PerplexityService: {str(e)}")
            raise

        # API 클라이언트 초기화
        try:
            self.anthropic = anthropic.Anthropic(
                api_key=os.getenv("XAI_API_KEY"), 
                base_url="https://api.x.ai"
            )
        except Exception as e:
            self.logger.error(f"Failed to initialize Anthropic client: {e}")
            self.anthropic = None
            
            # BERT 모델 초기화

            # Define search filters by research domain
        self.search_domain_filter = {
            "military": [
                "military", "defense", "security", "strategy", "tactics", "weapons", 
                "armed forces", "operations", "warfare", "combat", "intelligence"
            ],
            "technology": [
                "technology", "science", "R&D", "innovation", "systems", "equipment",
                "weapons systems", "engineering", "technical", "infrastructure", "capabilities"
            ],
            "policy": [
                "policy", "regulation", "legislation", "guidelines", "organization",
                "framework", "governance", "compliance", "standards", "protocols", "procedures"
            ],
            "education": [
                "education", "training", "development", "learning", "instruction",
                "curriculum", "doctrine", "methodology", "competency", "skills", "knowledge"
            ],
            "analysis": [
                "analysis", "evaluation", "validation", "measurement", "testing",
                "experimentation", "research", "assessment", "metrics", "verification"
            ],
            "operations": [
                "operations", "management", "maintenance", "support", "logistics",
                "administration", "service", "coordination", "implementation", "execution"
            ]
        }

        # Initial research steps template
        self.research_steps_template = [
            {
                "description": "Research Scope Definition and Objective Setting",
                "keywords": [
                    "research objectives", 
                    "scope definition", 
                    "requirements",
                    "project boundaries",
                    "success criteria"
                ],
                "methodology": "Literature review, Requirements analysis, Stakeholder consultation",
                "output_format": "Research Project Plan"
            },
            {
                "description": "Prior Research and Case Study Analysis",
                "keywords": [
                    "prior research",
                    "case studies",
                    "benchmarking",
                    "best practices",
                    "comparative analysis"
                ],
                "methodology": "Literature survey, Case study research, Comparative analysis",
                "output_format": "Analysis Report"
            },
            {
                "description": "Research Methodology Development",
                "keywords": [
                    "research methods",
                    "analytical approach",
                    "evaluation methods",
                    "research design",
                    "methodology framework"
                ],
                "methodology": "Methodology design, Framework development",
                "output_format": "Methodology Documentation"
            },
            {
                "description": "Data Collection and Analysis",
                "keywords": [
                    "data collection",
                    "data analysis",
                    "data processing",
                    "statistical analysis",
                    "findings synthesis"
                ],
                "methodology": "Data analysis, Statistical processing, Qualitative analysis",
                "output_format": "Analysis Results Report"
            }
        ]

    def init_bert_model(self):
        try:
        # 현재 단계 상태 초기화만 유지
            self.current_step = 1

            # 레이블 매핑 유지 (API 응답 처리에 필요)
            self.label_map = {
                0: "description",
                1: "keywords", 
                2: "methodology",
                3: "output_format"
            }
            
            self.logger.info("API-based classification initialized")
            
        except Exception as e:
            self.logger.error(f"Initialization failed: {str(e)}")
            raise

    def _generate_research_plan(self, project: Project) -> str:
        try:
            response = self.anthropic.messages.create(
                model="grok-beta",
                max_tokens=2000,
                temperature=0.7,
                messages=[{
                    "role": "user", 
                    "content": self._create_research_prompt(project)
                }]
            )
            # content 속성에서 텍스트 추출
            if hasattr(response, 'content'):
                if isinstance(response.content, list):
                    return response.content[0].text
                return response.content
            else:
                raise ValueError("API response does not contain expected content")

        except Exception as e:
            self.logger.error(f"Research plan generation failed: {str(e)}")
            raise HTTPException(status_code=500, detail=str(e))

    def _translate_to_korean(self, english_text: str) -> str:
        try:
            # 빈 텍스트 체크
            if not english_text.strip():
                return ""

            # 적절한 max_tokens 설정 (입력 텍스트 길이의 2배 정도로 설정)
            # 최소 1000 토큰, 최대 4000 토큰으로 제한
            estimated_tokens = min(max(len(english_text.split()) * 2, 1000), 4000)

            prompt = f"""
            Please translate the following English text to Korean. 
            This is a security/military research document, so please:

            1. Maintain technical accuracy in domain-specific terms
            2. Use standard Korean military/security terminology
            3. Preserve any technical acronyms in their original form
            4. Keep formatting and structure intact
            5. Ensure natural Korean language flow
            6. Maintain academic/formal tone

            Text to translate:
            {english_text}

            Translation guidelines:
            - Military terms should follow Korean military standard terminology
            - Technical terms should use established Korean translations when available
            - Preserve paragraph structure and formatting
            - Keep numbers, codes, and equations in their original form
            - Maintain citations and references in their original form
            - Use appropriate Korean honorifics for formal academic writing

            Please provide only the Korean translation without explanations or notes.
            """

            response = self.anthropic.messages.create(
                model="grok-beta",
                max_tokens=estimated_tokens,  # 수정된 부분
                temperature=0.1,
                system="You are an expert translator specializing in military and security domain translation from English to Korean. You have deep knowledge of both languages and domain-specific terminology.",
                messages=[{"role": "user", "content": prompt}]
            )

            translated_text = response.content[0].text.strip()
            return self._post_process_translation(translated_text)

        except Exception as e:
            self.logger.error(f"Translation failed: {str(e)}")
            self.logger.debug(f"Failed text: {english_text[:100]}...")
            raise HTTPException(status_code=500, detail=f"Translation failed: {str(e)}")

    def _post_process_translation(self, translated_text: str) -> str:
        """
        Post-process the translated text to ensure consistency and proper formatting.
        """
        try:
            # Preserve formatting characters
            preserved_chars = {
                '#': '###HASH###',
                '*': '###ASTERISK###',
                '_': '###UNDERSCORE###',
                '`': '###BACKTICK###',
                '[': '###LEFTBRACKET###',
                ']': '###RIGHTBRACKET###'
            }

            # Replace formatting characters with placeholders
            for char, placeholder in preserved_chars.items():
                translated_text = translated_text.replace(char, placeholder)

            # Clean up common translation artifacts
            cleaned_text = translated_text.replace('번역:', '').replace('번역 결과:', '')
            cleaned_text = cleaned_text.strip()

            # Restore formatting characters
            for char, placeholder in preserved_chars.items():
                cleaned_text = cleaned_text.replace(placeholder, char)

            # Ensure proper spacing around special characters
            cleaned_text = cleaned_text.replace('  ', ' ')
            cleaned_text = cleaned_text.replace(' :', ':')
            cleaned_text = cleaned_text.replace(' .', '.')
            cleaned_text = cleaned_text.replace(' ,', ',')

            return cleaned_text

        except Exception as e:
            self.logger.warning(f"Translation post-processing failed: {str(e)}")
            return translated_text

    def _batch_translate(self, texts: List[str], batch_size: int = 5) -> List[str]:
        """
        Translate multiple texts in batches to optimize API usage.
        """
        translated_texts = []
        
        try:
            for i in range(0, len(texts), batch_size):
                batch = texts[i:i + batch_size]
                
                # Create a combined prompt for the batch
                combined_text = "\n---SECTION BREAK---\n".join(batch)
                translated_batch = self._translate_to_korean(combined_text)
                
                # Split the translated text back into individual sections
                sections = translated_batch.split("\n---SECTION BREAK---\n")
                translated_texts.extend(sections)
                
        except Exception as e:
            self.logger.error(f"Batch translation failed: {str(e)}")
            raise
            
        return translated_texts

    def _translate_report_sections(self, english_report: str) -> str:
        """
        Translate report sections while preserving structure and formatting.
        Handles different section types appropriately.
        """
        try:
            # Split report into major sections
            sections = english_report.split('\n\n')
            translated_sections = []
            
            # Group sections for batch translation
            content_sections = []
            section_types = []  # Track if section is header or content
            
            for section in sections:
                if section.strip():
                    is_header = section.startswith('#') or section.strip().isupper()
                    if is_header:
                        # Split header and content
                        header_lines = [line for line in section.split('\n') if line.startswith('#') or line.strip().isupper()]
                        content_lines = [line for line in section.split('\n') if not (line.startswith('#') or line.strip().isupper())]
                        
                        if header_lines:
                            section_types.append('header')
                            content_sections.append('\n'.join(header_lines))
                        if content_lines:
                            section_types.append('content')
                            content_sections.append('\n'.join(content_lines))
                    else:
                        section_types.append('content')
                        content_sections.append(section)
            
            # Batch translate all sections
            translated_contents = self._batch_translate(content_sections)
            
            # Reconstruct the document
            current_section = []
            for section_type, translated_content in zip(section_types, translated_contents):
                if section_type == 'header':
                    if current_section:
                        translated_sections.append('\n'.join(current_section))
                        current_section = []
                    current_section.append(translated_content)
                else:
                    current_section.append(translated_content)
            
            if current_section:
                translated_sections.append('\n'.join(current_section))
            
            return '\n\n'.join(translated_sections)
            
        except Exception as e:
            self.logger.error(f"Report section translation failed: {str(e)}")
            return english_report
    def _create_research_prompt(self, project: Project) -> str:
        """Enhanced research prompt generation with 6-step structure"""
        return f"""
        You are an expert research project manager with extensive experience in analyzing and structuring research projects. Your expertise includes methodology development, data analysis, and comprehensive research planning.

        [Project Information]
        Title: {project.title}
        Description: {project.description}
        Evaluation Plan: {project.evaluation_plan}
        Submission Format: {project.submission_format}

        Please analyze this project and provide a detailed research plan following these steps:

        1. Project Analysis
        First, analyze the project thoroughly considering:
        - Core research objectives and scope
        - Key stakeholders and their requirements
        - Potential challenges and constraints
        - Available resources and limitations
        - Success criteria and evaluation metrics

        Format your analysis as:
        [Project Analysis]
        Core Objectives:
        - [List key objectives]
        Stakeholders:
        - [List stakeholders and their needs]
        Challenges:
        - [List potential challenges]
        Resources:
        - [List available resources]
        Success Metrics:
        - [List key metrics]

        2. Research Stage Planning
        For each research stage:
        - Define clear objectives
        - Identify required resources
        - Specify methodologies
        - Determine deliverables
        - Establish success criteria

        3. Literature Review Planning
        For each stage:
        - List key search terms and phrases
        - Define review scope
        - Specify source types
        - Outline review methodology
        - Plan synthesis approach

        4. Methodology Development
        For each stage:
        - Detail research methods
        - Outline data collection
        - Specify analysis techniques
        - Define quality controls
        - Identify validation methods

        5. Risk Assessment
        Evaluate:
        - Technical risks
        - Resource constraints
        - Timeline challenges
        - Quality concerns
        - External dependencies

        6. Research Output Structure
        Present each research stage using this format:

        ## [Stage Number] Stage
        **Description**: [Detailed description]
        **Keywords**: [Research terms]
        **Research Method**: [Methodology]
        **Expected Output**: [Deliverables]
        **Success Criteria**: [Measurable outcomes]
        **Risk Mitigation**: [Specific measures]

        Quality Guidelines:
        - Ensure clear, measurable objectives
        - Use appropriate methodologies
        - Include specific deliverables
        - Address identified risks
        - Align with project goals
        - Meet stakeholder requirements

        Provide your complete response with:
        1. Project Analysis Summary
        2. Detailed Research Stages
        3. Risk Assessment Matrix
        4. Resource Allocation Plan
        5. Timeline Overview
        6. Quality Control Measures
        """

    def _parse_research_plan(self, plan_text: str) -> Dict:
        try:
            steps = []
            current_step = None
                
            lines = [line.strip() for line in plan_text.split('\n') if line.strip()]
            
            for line in lines:
                # Stage 시작 감지
                if "Stage" in line and ("**" in line or "##" in line):
                    if current_step:
                        steps.append(current_step)
                    current_step = {
                        "description": "",
                        "keywords": [],
                        "methodology": "",
                        "output_format": ""  # expected_output 대신 output_format 사용
                    }
                    continue
                    
                if not current_step:
                    continue
                    
                # 콘텐츠 파싱
                if "**Description**:" in line or "Description:" in line:
                    current_step["description"] = line.split(":", 1)[1].strip()
                elif "**Keywords**:" in line or "Keywords:" in line:
                    keywords_text = line.split(":", 1)[1].strip()
                    keywords = [k.strip() for k in keywords_text.split(",")]
                    current_step["keywords"] = keywords
                elif "**Research Method**:" in line or "Research Method:" in line:
                    current_step["methodology"] = line.split(":", 1)[1].strip()
                # Expected Output을 output_format으로 변환
                elif "**Expected Output**:" in line or "Expected Output:" in line:
                    current_step["output_format"] = line.split(":", 1)[1].strip()
            
            # 마지막 단계 추가
            if current_step:
                steps.append(current_step)
                
            if not steps:
                self.logger.warning("No steps were parsed from the plan text")
                raise ValueError("Could not parse any valid research steps from the AI response")
                
            return {
                "research_steps": steps,
                "project_analysis": self._extract_project_analysis(plan_text),
                "risk_assessment": self._extract_risk_assessment(plan_text),
                "resource_plan": {},
                "timeline": {},
                "quality_measures": {}
            }
            
        except Exception as e:
            self.logger.error(f"Research plan parsing failed: {str(e)}")
            raise

    def _extract_project_analysis(self, plan_text: str) -> Dict:
        """Project Analysis 섹션에서 정보 추출"""
        analysis = {
            "core_objectives": [],
            "stakeholders": [],
            "challenges": [],
            "resources": [],
            "success_metrics": []
        }
        
        try:
            # Project Analysis 섹션 찾기
            if "Project Analysis" in plan_text:
                section = plan_text.split("Project Analysis")[1].split("###")[0]
                
                # 각 하위 섹션 파싱
                for line in section.split("\n"):
                    line = line.strip()
                    if line.startswith("- "):
                        if "Core Objectives" in section:
                            analysis["core_objectives"].append(line[2:])
                        elif "Stakeholders" in section:
                            analysis["stakeholders"].append(line[2:])
                        elif "Challenges" in section:
                            analysis["challenges"].append(line[2:])
                        elif "Resources" in section:
                            analysis["resources"].append(line[2:])
                        elif "Success Metrics" in section:
                            analysis["success_metrics"].append(line[2:])
        except Exception as e:
            self.logger.error(f"Error extracting project analysis: {str(e)}")
        
        return analysis

    def _extract_risk_assessment(self, plan_text: str) -> Dict:
        """Risk Assessment 섹션에서 정보 추출"""
        risks = []
        
        try:
            if "Risk Assessment" in plan_text:
                section = plan_text.split("Risk Assessment")[1].split("###")[0]
                for line in section.split("\n"):
                    line = line.strip()
                    if line.startswith("| ") and "Risk" in line:
                        parts = line.split("|")
                        if len(parts) >= 5:
                            risks.append({
                                "category": parts[1].strip(),
                                "risk": parts[2].strip(),
                                "likelihood": parts[3].strip(),
                                "impact": parts[4].strip(),
                                "mitigation": parts[5].strip() if len(parts) > 5 else ""
                            })
        except Exception as e:
            self.logger.error(f"Error extracting risk assessment: {str(e)}")
        
        return {"risks": risks}
    
    def _classify_section(self, text: str) -> Optional[str]:
        if not text.strip():
            return None
        
        # Section keywords for initial filtering
        section_keywords = {
            "description": ["description", "step description", "**description**"],
            "keywords": ["keyword", "keywords", "**keyword**", "**keywords**"],
            "methodology": ["method", "methodology", "research method", "**method**"],
            "expected_output": ["output", "expected output", "deliverables", "**expected output**"],
            "success_criteria": ["success criteria", "success metrics", "**success criteria**"],
            "risk_mitigation": ["risk mitigation", "risks", "**risk mitigation**"]
        }

        # Text preprocessing
        text = text.lower().strip()

        # Initial keyword-based classification
        for section, keywords in section_keywords.items():
            if any(keyword.lower() in text for keyword in keywords):
                return section

        # Prepare sentences for comparison
        sentences_to_compare = list(section_keywords.keys())
        
        try:
            # Prepare payload for API call
            payload = {
                "inputs": {
                    "source_sentence": text,
                    "sentences": sentences_to_compare
                }
            }
            def query(payload):
                response = requests.post(API_URL, headers=headers, json=payload)
                return response.json()
            # Make API call
            output = query(payload)
            
            # Check if the API returned embeddings
            if 'embeddings' in output:
                # Assuming the API returns similarities or you calculate them from embeddings
                # For simplicity, let's assume it directly gives similarities:
                similarities = output['embeddings']
                max_similarity_index = similarities.index(max(similarities))
                predicted_category = sentences_to_compare[max_similarity_index]
                return predicted_category
            else:
                self.logger.warning(f"Unexpected API response format: {output}")
                return "description"  # Default fallback if we can't interpret the response
        
        except Exception as e:
            self.logger.warning(f"API call failed: {str(e)}")
            return "description"
    def create_project(self, user_id: int, project_data: Dict) -> Project:
        """Enhanced project creation with new structure"""
        try:
            transaction = self.db.begin()
            with transaction:
                new_project = Project(
                    user_id=user_id,
                    title=project_data['title'],
                    description=project_data['description'],
                    evaluation_plan=project_data.get('evaluation_plan', ''),
                    submission_format=project_data.get('submission_format', ''),
                    metadata1=project_data.get('metadata1', {}),
                    research_field=project_data.get('research_field', '안보'),
                    evaluation_status='pending',
                    total_steps=0,
                    completed_steps=0
                )
                self.db.add(new_project)
                self.db.flush()
                
                # Generate enhanced research plan
                plan_result = self._generate_research_plan(new_project)
                parsed_plan = self._parse_research_plan(plan_result)
                
                # Create research steps with enhanced structure
                for step_number, step_data in enumerate(parsed_plan['research_steps'], 1):
                    step = ResearchStep(
                        project_id=new_project.id,
                        step_number=step_number,
                        description=step_data['description'],
                        keywords=step_data.get('keywords', []),
                        methodology=step_data.get('methodology', ''),
                        output_format=step_data.get('output_format', ''),
                        status='pending',
                        progress_percentage=0
                    )
                    self.db.add(step)
                    
                # Store additional analysis data in project metadata
                new_project.metadata1.update({
                    'project_analysis': parsed_plan['project_analysis'],
                    'risk_assessment': parsed_plan['risk_assessment'],
                    'resource_plan': parsed_plan['resource_plan'],
                    'timeline': parsed_plan['timeline'],
                    'quality_measures': parsed_plan['quality_measures']
                })
                
                new_project.total_steps = len(parsed_plan['research_steps'])
                self.db.commit()            
                return new_project
                
        except Exception as e:
            self.logger.error(f"Project creation failed: {str(e)}")
            raise HTTPException(status_code=500, detail=str(e))

    def _validate_steps(self, steps: List[Dict]) -> List[Dict]:
        for step in steps:
            if not all(k in step for k in ["description", "keywords", "methodology", "output_format"]):
                raise ValueError("필수 섹션이 누락되었습니다.")
        return steps


    def _save_research_plan(self, project_id: int, steps_data: List[Dict]) -> None:
        """연구 계획 저장"""
        try:
            df = pd.DataFrame(steps_data)
            df.to_csv(f"research_plans/project_{project_id}_plan.csv", index=False)
        except Exception as e:
            self.logger.error(f"Failed to save research plan: {str(e)}")
            # 계획 저장 실패는 크리티컬하지 않으므로 예외를 전파하지 않음

    def _extract_research_keywords(self, step: ResearchStep, project: Project) -> List[str]:
        """
        Extract relevant keywords for literature search from research step description.
        Uses advanced prompt engineering and domain-specific filtering.
        """
        try:
            # Create enhanced prompt with clear instructions and structure
            prompt = f"""
            Please analyze the following research project information and extract key academic search terms.
            Focus on identifying specific, technical, and domain-relevant keywords that would be effective
            for academic database searches.

            [Project Information]
            Title: {project.title}
            Description: {project.description}
            Research Field: {project.research_field}

            [Current Research Stage]
            Stage Number: {step.step_number}
            Description: {step.description}
            Methodology: {step.methodology}

            Please provide keywords following these criteria:

            1. Technical Terms:
            - Domain-specific terminology
            - Technical concepts and methodologies
            - Specialized equipment or systems

            2. Research Concepts:
            - Theoretical frameworks
            - Research methodologies
            - Analytical approaches

            3. Application Areas:
            - Specific use cases
            - Implementation contexts
            - practical applications

            4. Related Fields:
            - Interdisciplinary connections
            - Related research domains
            - Relevant subfields

            Format your response as a structured list with categories:
            [Technical]
            - keyword1
            - keyword2

            [Research]
            - keyword3
            - keyword4

            [Application]
            - keyword5
            - keyword6

            [Related]
            - keyword7
            - keyword8

            Guidelines:
            - Prefer specific terms over general ones
            - Include both broader and narrower terms
            - Consider standard academic terminology
            - Include relevant acronyms and full terms
            - Focus on searchable academic keywords
            """

            # Enhanced API call with better parameters
            response = self.anthropic.messages.create(
                model="grok-beta",
                max_tokens=1000,
                temperature=0.1,
                system="""You are an expert research librarian specializing in academic keyword extraction 
                        and literature search optimization. Focus on identifying precise, relevant, and 
                        academically appropriate search terms.""",
                messages=[{"role": "user", "content": prompt}]
            )

            # Enhanced keyword extraction and processing
            raw_keywords = []
            current_category = None
            
            # Parse structured response
            for line in response.content[0].text.split('\n'):
                line = line.strip()
                if line.startswith('[') and line.endswith(']'):
                    current_category = line[1:-1]
                    continue
                if line.startswith('- '):
                    keyword = line[2:].strip()
                    if keyword:
                        raw_keywords.append(keyword)

            # Process and clean keywords
            processed_keywords = self._process_keywords(raw_keywords)
            
            # Combine with existing keywords
            combined_keywords = list(set(processed_keywords + step.keywords))
            
            # Apply domain-specific filtering with weighted relevance
            filtered_keywords = self._apply_domain_filters(
                combined_keywords,
                project.research_field,
                step.step_number
            )
            
            return filtered_keywords

        except Exception as e:
            self.logger.error(f"Keyword extraction failed: {str(e)}")
            self.logger.debug(f"Error details: {str(e)}", exc_info=True)
            return step.keywords

    def _process_keywords(self, keywords: List[str]) -> List[str]:
        """
        Process and clean extracted keywords.
        """
        processed = []
        for keyword in keywords:
            # Remove special characters and normalize spacing
            cleaned = re.sub(r'[^\w\s-]', ' ', keyword)
            cleaned = ' '.join(cleaned.split())
            
            # Handle acronyms
            if cleaned.isupper() and len(cleaned) <= 5:
                processed.append(cleaned)  # Keep acronyms as-is
            else:
                # Convert to lowercase for standard terms
                processed.append(cleaned.lower())
                
                # Add acronym if it's a multi-word term
                words = cleaned.split()
                if len(words) > 1:
                    acronym = ''.join(word[0].upper() for word in words)
                    if len(acronym) >= 2:
                        processed.append(acronym)
        
        return list(set(processed))  # Remove duplicates

    def _apply_domain_filters(self, keywords: List[str], research_field: str, step_number: int) -> List[str]:
        """
        Apply domain-specific filtering with weighted relevance scoring.
        """
        field_keywords = self.search_domain_filter.get(research_field, [])
        relevance_scores = {}
        
        for keyword in keywords:
            score = 0
            # Check direct matches with field keywords
            for field_term in field_keywords:
                if field_term in keyword.lower():
                    score += 2
                    
            # Consider keyword length (prefer more specific terms)
            if len(keyword.split()) > 1:
                score += 1
                
            # Consider step number (early steps may need broader terms)
            if step_number <= 2:
                score += 1  # Boost broader terms for early steps
            
            # Store score
            relevance_scores[keyword] = score
        
        # Filter keywords based on minimum relevance score
        min_score = 1 if step_number <= 2 else 2
        filtered_keywords = [k for k, score in relevance_scores.items() if score >= min_score]
        
        # Sort by relevance score
        return sorted(filtered_keywords, key=lambda k: relevance_scores[k], reverse=True)



    async def _generate_final_report(self, project: Project, results: List[Dict]) -> Dict:
        """Generate comprehensive final research report using Claude API"""
        try:
            # Ensure that the Anthropic client is initialized
            if self.anthropic is None:
                raise Exception("Anthropic client is not initialized. Please check the API key or base URL.")

            # Convert step results to formatted string
            steps_results = "\n\n".join([
                f"[Stage {i+1} Results]\n{result['result']}"
                for i, result in enumerate(results)
            ])
            prompt = f"""
                        You are a highly skilled research analyst specializing in military and security projects. Your task is to synthesize research results into a comprehensive, detailed, and professional final report. Please carefully review the following project information and research results:

<project_title>
{project.title}
</project_title>

<project_description>
{project.description}
</project_description>

<project_evaluation_plan>
{project.evaluation_plan}
</project_evaluation_plan>

<steps_results>
{steps_results}
</steps_results>

Your goal is to create a final report that adheres to the following format:

<submission_format>
{project.submission_format}
</submission_format>

Before writing the final report, please conduct a thorough analysis of the provided information in <research_breakdown> tags to show your thought process and ensure a comprehensive understanding of the project and its results.

Instructions for analysis:
1. Carefully review the project title, description, and evaluation plan.
2. Examine the stage-wise research results in detail.
3. Identify key findings, trends, and patterns in the research.
4. Consider the practical applications of the research in a military/security context.
5. Determine actionable insights that would be valuable to stakeholders.
6. Assess any security implications or sensitivities in the research.
7. Plan how to structure the information to ensure a logical flow in the final report.
8. Consider how to maintain research consistency and completeness throughout the report.
9. Summarize key points from each section of the project information.
10. Identify and list potential challenges or limitations in the research.
11. Outline a preliminary structure for the final report.

After completing your analysis, write the final report according to the specified format. Ensure that your report is as detailed and professional as possible, demonstrating a high level of expertise in the subject matter.

Guidelines for the final report:
1. Maintain a logical flow between sections, ensuring that each part of the report builds upon the previous information.
2. Emphasize practical applications in the military/security context, clearly explaining how the research findings can be implemented or utilized.
3. Provide specific, actionable insights for stakeholders, detailing how they can leverage the research results.
4. Address security implications and sensitivities with appropriate caution and discretion.
5. Use precise, technical language where appropriate, but ensure clarity for a professional audience.
6. Include relevant data, statistics, and examples to support your conclusions and recommendations.
7. Be as comprehensive as possible, exploring all aspects of the research in depth.
8. Conclude with a strong summary that reinforces the key findings and their significance.

Remember, the quality and depth of your report are crucial. Take the time to craft a document that truly reflects the importance and complexity of the research.

Please begin your response with your analysis, followed by the final report.
            """
            
            # 동기적으로 API 호출
            response = self.anthropic.messages.create(
                model="grok-beta",
                max_tokens=4000,
                temperature=0.3,
                system="You are a senior security research expert crafting a comprehensive final report with deep understanding of military and security domains.",
                messages=[{"role": "user", "content": prompt}]
            )
            english_report = response.content[0].text
            
            # 비동기 작업을 동기적으로 처리
            korean_report = self._translate_report_sections(english_report)
            
            # 최종 보고서 저장
            project.metadata1 = project.metadata1 or {}
            project.metadata1['final_reports'] = {
                'final_report_en': english_report,
                'final_report_kr': korean_report,
                'generated_at': datetime.utcnow().isoformat()
            }
            project.evaluation_status = "completed"
            
            # English report to DOCX
            self._save_report_to_docx(english_report, f"final_report_en_{project.id}.docx")
            
            # Korean report to DOCX
            self._save_report_to_docx(korean_report, f"final_report_kr_{project.id}.docx")
            
            self.db.commit()

            return {
                "project_id": project.id,
                "title": project.title,
                "final_report_en": english_report,
                "final_report_kr": korean_report,
                "step_results": results,
                "generated_at": datetime.utcnow().isoformat(),
                "status": "completed"
            }

        except Exception as e:
            self.logger.error(f"Final report generation failed: {str(e)}")
            raise

    def _save_report_to_docx(self, report_text: str, filename: str):
        """Save report text to a DOCX file."""
        try:
            doc = Document()
            doc.add_heading('Final Research Report', level=1)
            
            # Add content
            for section in report_text.split('\n\n'):
                if section.startswith('#') or section.strip().isupper():
                    # Section header
                    doc.add_heading(section, level=2)
                else:
                    # Regular content
                    doc.add_paragraph(section)
            
            # Save the document
            doc.save(f"reports/{filename}")
            
        except Exception as e:
            self.logger.error(f"Failed to save report as DOCX: {str(e)}")
            raise


    def _generate_step_prompt(self, step: ResearchStep, project: Project, literature: List[Dict]) -> str:
        """Generate stage-specific research prompts including prior research"""
        base_prompt = f"""
        [Project Information]
        Title: {project.title}
        Description: {project.description}

        [Current Stage Information]
        Stage: Stage {step.step_number}
        Description: {step.description}
        Methodology: {step.methodology}
        Expected Deliverable: {step.output_format}
        
        [Relevant Prior Research]
        Please consider the following prior research in your analysis:
        """
        
        # Add prior research information
        for idx, source in enumerate(literature, 1):
            base_prompt += f"""
            Source {idx}:
            Title: {source['metadata']['title']}
            URL: {source['metadata']['url']}
            Key Content: {source['pageContent']}
            """
        
        # Add base requirements template
        base_prompt += """
        Please conduct your research according to these requirements:
        1. Strictly adhere to the specified methodology
        2. Conduct objective and logical analysis
        3. Present concrete and practical results
        4. Consider the unique aspects of military/security domain
        5. Actively incorporate insights from prior research
        """
        
        # Add stage-specific prompts
        if step.step_number == 1:
            base_prompt += """
            For this initial stage, please focus on:
            1. Research background and justification
            2. Clear research objectives
            3. Specific research scope definition
            4. Analysis of potential constraints
            5. Core research questions identification

            Structure your response as follows:
            - Research Background
            - Research Objectives
            - Research Scope
            - Constraints
            - Research Questions
            """
        elif step.step_number == 2:
            base_prompt += """
            For the data collection and analysis stage, include:
            1. Types and scope of data to be collected
            2. Data collection methodology
            3. Analysis framework
            4. Key analysis metrics
            5. Expected analysis outcomes

            Structure your response as follows:
            - Data Collection Plan
            - Analysis Methodology
            - Key Analysis Metrics
            - Preliminary Analysis Results
            - Implications
            """
        elif step.step_number == 3:
            base_prompt += """
            For the methodology development stage, include:
            1. Research approach selection
            2. Detailed methodology design
            3. Validation approach
            4. Risk factor analysis
            5. Mitigation strategy development

            Structure your response as follows:
            - Research Approach
            - Detailed Methodology
            - Validation Plan
            - Risk Analysis
            - Mitigation Strategies
            """
        elif step.step_number == 4:
            base_prompt += """
            For the results analysis and validation stage, focus on:
            1. Comprehensive data analysis
            2. Hypothesis validation
            3. Interpretation of results
            4. Implications identification
            5. Recommendations development

            Structure your response as follows:
            - Comprehensive Analysis Results
            - Hypothesis Validation Results
            - Key Findings
            - Policy Implications
            - Recommendations
            """

        return base_prompt
    
    async def _fetch_literature(self, keywords: List[str]) -> List[Dict]:
        """Fetch literature using Perplexity AI API with SSL verification disabled"""
        literature = []
        
        try:
            # SSL 검증을 비활성화한 컨텍스트 생성 추가
            ssl_context = ssl.create_default_context()
            ssl_context.check_hostname = False
            ssl_context.verify_mode = ssl.CERT_NONE
            
            connector = aiohttp.TCPConnector(ssl=ssl_context)
            
            async with aiohttp.ClientSession(connector=connector) as session:
                for keyword in keywords:
                    try:
                        query = f"Academic references on {keyword}"
                        config = SearchConfig(
                            search_domain="academic",
                            search_recency_filter="year",
                            temperature=0.2
                        )
                        
                        result = await self.perplexity_service.search_references(
                            query=query,
                            model="llama-3.1-sonar-small-128k-online",
                            config=config
                        )
                        
                        if result and 'references' in result:
                            literature.extend(result['references'])
                        
                        await asyncio.sleep(0.5)
                        
                    except Exception as e:
                        self.logger.error(f"Error fetching literature for keyword '{keyword}': {str(e)}")
                        continue
            
            # 중복 제거
            seen_urls = set()
            unique_literature = []
            for item in literature:
                url = item.get('url', '')
                if url and url not in seen_urls:
                    seen_urls.add(url)
                    unique_literature.append(item)
            
            return unique_literature
            
        except Exception as e:
            self.logger.error(f"Literature fetch failed: {str(e)}")
            return []
        
        finally:
            if 'connector' in locals():
                await connector.close()

    async def _execute_step(self, step: ResearchStep, project: Project) -> Dict:
        """단일 연구 단계 실행"""
        try:
            # 1. 키워드 추출
            keywords = self._extract_research_keywords(step, project)
            
            # 2. 선행 연구 수집 (비동기 호출)
            literature = await self._fetch_literature(keywords)
            
            # 3. 연구 프롬프트 생성
            prompt = self._generate_step_prompt(step, project, literature)
            
            # 4. Claude API로 연구 수행
            response = self.anthropic.messages.create(
                model="grok-beta",
                max_tokens=2000,
                temperature=0.3,
                system="You are a security research expert conducting a detailed analysis with expertise in military and security domains.",
                messages=[{"role": "user", "content": prompt}]
            )
            
            # 5. 결과 구조화
            return {
                "step_number": step.step_number,
                "result": response.content[0].text,
                "keywords": keywords,
                "literature": [
                    {
                        "title": ref.get("title", ""),
                        "authors": ref.get("authors", []),
                        "abstract": ref.get("abstract", ""),
                        "url": ref.get("url", ""),
                        "publication_date": ref.get("publication_date", "")
                    }
                    for ref in literature
                ],
                "timestamp": datetime.utcnow().isoformat()
            }
            
        except Exception as e:
            self.logger.error(f"Step execution failed: {str(e)}")
            raise

    async def execute_all_steps(self, project_id: int) -> Dict:
        """모든 연구 단계 실행 with improved error handling"""
        project = self.db.query(Project).get(project_id)
        if not project:
            raise HTTPException(status_code=404, detail="Project not found")
        
        try:
            steps = self.db.query(ResearchStep).filter_by(
                project_id=project_id
            ).order_by(ResearchStep.step_number).all()
            
            results = []
            
            for step in steps:
                try:
                    # 단계 시작
                    step.status = 'in_progress'
                    step.started_at = datetime.utcnow()
                    self.db.commit()
                    
                    # 각 단계별 타임아웃 설정
                    async with asyncio.timeout(300):  # 5분 타임아웃
                        result = await self._execute_step(step, project)
                    
                    # 성공 처리
                    step.status = 'completed'
                    step.completed_at = datetime.utcnow()
                    step.result = result
                    project.completed_steps += 1
                    results.append(result)
                    
                except asyncio.TimeoutError:
                    step.status = 'failed'
                    step.error_message = "Step execution timed out"
                    self.logger.error(f"Step {step.step_number} timed out")
                    
                except Exception as e:
                    step.status = 'failed'
                    step.error_message = str(e)
                    self.logger.error(f"Step {step.step_number} execution failed: {e}")
                
                finally:
                    self.db.commit()
            
            # 최종 보고서 생성
            final_result = await self._generate_final_report(project, results)
            return final_result
            
        except Exception as e:
            project.evaluation_status = 'failed'
            self.db.commit()
            self.logger.error(f"Project execution failed: {str(e)}")
            raise
    def _translate_report_sections(self, english_report: str) -> str:
        """Translate report sections while preserving structure"""
        try:
            # Split report into sections
            sections = english_report.split('\n\n')
            translated_sections = []
            
            for section in sections:
                # Preserve section headers
                if section.startswith('#') or section.strip().isupper():
                    header = section.split('\n')[0]
                    content = '\n'.join(section.split('\n')[1:])
                    translated_content = self._translate_to_korean(content)
                    translated_sections.append(f"{header}\n{translated_content}")
                else:
                    translated_sections.append(self._translate_to_korean(section))
            
            return '\n\n'.join(translated_sections)
            
        except Exception as e:
            self.logger.error(f"Report translation failed: {str(e)}")
            return english_report


    def _extract_keywords(self, text: str) -> List[str]:
        """키워드 추출"""
        import torch
        
        # 키워드 섹션 헤더 제거
        for header in ["keyword:", "keyword", "**keyword**:", "**keyword**"]:
            text = text.replace(header, "")
        
        try:
            # NER 모델을 사용한 keyword 추출
            inputs = self.ner_tokenizer(
                text,
                padding=True,
                truncation=True,
                max_length=128,
                return_tensors="pt"
            ).to(self.device)
            
            with torch.no_grad():
                outputs = self.ner_model(**inputs)
                predictions = torch.argmax(outputs.logits, dim=2)
            
            # 토큰을 키워드로 변환
            tokens = self.ner_tokenizer.convert_ids_to_tokens(inputs["input_ids"][0])
            keywords = []
            current_keyword = []
            
            for token, pred in zip(tokens, predictions[0]):
                if pred == 1:  # B-KW
                    if current_keyword:
                        keywords.append("".join(current_keyword))
                        current_keyword = []
                    current_keyword.append(token)
                elif pred == 2:  # I-KW
                    current_keyword.append(token)
            
            if current_keyword:
                keywords.append("".join(current_keyword))
            
            # 후처리: 특수 토큰 제거 및 정제
            keywords = [
                kw.replace("#", "").replace("##", "").strip()
                for kw in keywords
                if kw not in ["[CLS]", "[SEP]", "[PAD]"]
            ]
            
            return [kw for kw in keywords if kw]
            
        except Exception as e:
            # NER 모델 실패 시 기본 키워드 추출 로직 사용
            self.logger.warning(f"NER keyword extraction failed: {str(e)}")
            keywords = []
            for keyword in text.split(","):
                cleaned = keyword.strip().strip("[]").strip()
                if cleaned:
                    keywords.append(cleaned)
            return keywords

    def _process_section_content(self, text: str) -> str:
        """섹션 내용 처리"""
        import re
        # 특수 문자 및 마크다운 제거
        text = re.sub(r'\*\*|\[|\]|\#', '', text)
        return text.strip()

    def _post_process_steps(self, steps: List[Dict]) -> List[Dict]:
        """단계 데이터 후처리"""
        processed_steps = []
        for step in steps:
            # 키워드 중복 제거 및 정제
            step["keywords"] = list(set(step["keywords"]))
            
            # 빈 필드 기본값 설정
            if not step["description"]:
                step["description"] = f"{step['step_number']}단계"
            if not step["methodology"]:
                step["methodology"] = "방법론 미정의"
            if not step["output_format"]:
                step["output_format"] = "결과물 미정의"
            
            processed_steps.append(step)
        
        return processed_steps
    def analyze_text_similarity(self, text1: str, text2: str) -> float:
        """텍스트 유사도 분석 - Voyage AI 사용"""
        try:
            result = self.voyage.embed([text1, text2], model="voyage-3-lite", input_type="document")
            embeddings = result.embeddings
            
            similarity = np.dot(embeddings[0], embeddings[1]) / (
                np.linalg.norm(embeddings[0]) * np.linalg.norm(embeddings[1])
            )
            
            return float(similarity)
        except Exception as e:
            self.logger.error(f"Text similarity analysis failed: {str(e)}")
            raise